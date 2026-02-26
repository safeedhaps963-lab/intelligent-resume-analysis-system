"""
app/utils/pdf_parser.py - PDF and Document Parser Utility
==========================================================
This module provides utilities for parsing resume files
in various formats (PDF, DOCX, TXT).

Supports:
- PDF files (using PyPDF2)
- DOCX files (using python-docx)
- TXT files (plain text)

Each parser extracts text content while preserving
structure as much as possible.
"""

import os
import re
from typing import Optional, Tuple


def parse_resume_file(file_path: str) -> str:
    """
    Parse a resume file and extract text content.
    
    Automatically detects file type based on extension
    and uses the appropriate parser.
    
    Args:
        file_path: Path to the resume file
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    
    Example:
        text = parse_resume_file('/uploads/resume.pdf')
        print(f"Extracted {len(text)} characters")
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Parse based on extension
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    elif ext == '.doc':
        return parse_doc(file_path)
    elif ext == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def parse_pdf(file_path: str) -> str:
    """
    Parse PDF file and extract text.
    
    Uses PyPDF2 library to read PDF pages and extract text.
    Handles multi-page documents and attempts to preserve
    paragraph structure.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        str: Extracted text content
    
    Example:
        text = parse_pdf('/uploads/resume.pdf')
    """
    try:
        # 1. New Primary Engine: PyMuPDF (fitz) - Extremely robust and fast
        try:
            import fitz
            doc = fitz.open(file_path)
            text_parts = []
            has_images = False
            
            for page in doc:
                # Get text blocks
                blocks = page.get_text("blocks")
                # block structure: (x0, y0, x1, y1, "text", block_no, block_type)
                
                if not blocks:
                    continue

                # Sort blocks by x coordinate first to help identify columns, 
                # then by y coordinate.
                # However, a better way for resumes is to group by column or use a heuristic.
                # Let's use a "Column-Aware Sort"
                # We sort primarily by y, but allows for a "tolerance" in y 
                # to keep things in the same line, then sort by x.
                # BUT for multi-column, we want to read all of column 1 then all of column 2.
                
                # Heuristic: If there's a large gap in X and the vertical ranges overlap significantly, 
                # it's likely multi-column.
                
                # Use PyMuPDF's default "text" extraction which is actually quite good at 
                # following the reading order layout. "blocks" with manual sort (y, x) 
                # is what causes interleaving.
                
                page_text = page.get_text("text", sort=True)
                
                if page_text:
                    # Basic normalization
                    page_text = page_text.replace('\u0000', '') # Remove null bytes
                    page_text = re.sub(r'[\u200b-\u200d\ufeff]', '', page_text) # Remove zero-width spaces
                    text_parts.append(page_text)
                
                if len(page.get_images()) > 0:
                    has_images = True
            doc.close()
            
            full_text = '\n\n'.join(text_parts)
            if full_text and len(full_text.strip()) > 30:
                print(f"DEBUG: Success extracting with PyMuPDF (Default Sort) ({len(full_text)} chars)")
                return _clean_extracted_text(full_text)
            
            # If no text but has images, it's a scanned PDF
            if not full_text.strip() and has_images:
                print(f"DEBUG: Detected image-only PDF: {file_path}")
                return "__IMAGE_ONLY_PDF__"
                
        except Exception as fitz_err:
            print(f"DEBUG: PyMuPDF failed: {str(fitz_err)}")

        # ... (rest of fallbacks)
        # 2. Secondary Engine: pdfminer.six
        try:
            from pdfminer.high_level import extract_text as miner_extract
            full_text = miner_extract(file_path)
            if full_text and len(full_text.strip()) > 50:
                print(f"DEBUG: Success extracting with pdfminer.six ({len(full_text)} chars)")
                return _clean_extracted_text(full_text)
        except Exception as miner_err:
            print(f"DEBUG: pdfminer.six failed: {str(miner_err)}")

        # 3. Third Engine: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text_parts.append(t)
                
                full_text = '\n\n'.join(text_parts)
                if full_text and len(full_text.strip()) > 50:
                    print(f"DEBUG: Success extracting with pdfplumber ({len(full_text)} chars)")
                    return _clean_extracted_text(full_text)
        except Exception as plumber_err:
            print(f"DEBUG: pdfplumber failed: {str(plumber_err)}")

        # 4. Fourth Engine: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                t = page.extract_text(extraction_mode="layout")
                if not t: t = page.extract_text()
                if t: text_parts.append(t)
            
            full_text = '\n\n'.join(text_parts)
            if full_text and len(full_text.strip()) > 20:
                print(f"DEBUG: Success extracting with pypdf ({len(full_text)} chars)")
                return _clean_extracted_text(full_text)
        except Exception as pypdf_err:
            print(f"DEBUG: pypdf failed: {str(pypdf_err)}")

        print(f"DEBUG: All PDF extraction engines failed for {file_path}")
        return ""
    
    except Exception as e:
        print(f"DEBUG: Critical error in parse_pdf {file_path}: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: str) -> str:
    """
    Parse DOCX file and extract text.
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_parts)
        return _clean_extracted_text(full_text)
    except Exception as e:
        print(f"DEBUG: Failed to parse DOCX {file_path}: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_doc(file_path: str) -> str:
    """
    Parse legacy DOC file.
    """
    try:
        import subprocess
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
        if result.returncode == 0:
            return _clean_extracted_text(result.stdout)
        else:
            raise ValueError("antiword failed")
    except Exception as e:
        raise ValueError(f"Legcy DOC parsing failed. Please convert to DOCX or PDF. Detail: {str(e)}")


def parse_txt(file_path: str) -> str:
    """
    Parse plain text file with multiple encoding attempts.
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                content = f.read()
                if content:
                    return _clean_extracted_text(content)
        except Exception:
            continue
    raise ValueError("Could not read text file with any standard encoding")


def _clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text.
    
    Performs various cleaning operations:
    - Remove excessive whitespace
    - Fix line breaks
    - Remove special characters
    - Normalize unicode
    
    Args:
        text: Raw extracted text
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    # Replace non-breaking spaces with standard spaces
    text = text.replace('\xa0', ' ')
    
    # Replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Replace multiple newlines with double newline
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove leading/trailing whitespace from lines
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove only problematic control characters (like null bytes) but keep others
    # and all unicode. This is safer than a whitelist.
    bad_chars = '\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f'
    text = "".join(ch for ch in text if ch not in bad_chars)
    
    return text.strip()


def get_file_info(file_path: str) -> dict:
    """
    Get information about a file.
    
    Args:
        file_path: Path to the file
    
    Returns:
        dict: File information (size, type, etc.)
    """
    if not os.path.exists(file_path):
        return {}
    
    stat = os.stat(file_path)
    _, ext = os.path.splitext(file_path)
    
    return {
        'size': stat.st_size,
        'extension': ext.lower(),
        'filename': os.path.basename(file_path),
        'modified': stat.st_mtime
    }


def estimate_page_count(text: str) -> int:
    """
    Estimate the number of pages based on text length.
    
    Uses average words per page (approximately 400-500 words).
    
    Args:
        text: Resume text content
    
    Returns:
        int: Estimated page count
    """
    word_count = len(text.split())
    
    # Assume ~450 words per page average
    pages = max(1, round(word_count / 450))
    
    return pages